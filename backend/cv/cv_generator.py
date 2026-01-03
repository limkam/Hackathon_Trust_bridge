"""
CV Generator Module
Handles CV generation with AI enhancements, ATS optimization, and PDF/Word export.
"""
from typing import Dict, Any, Optional, List
from sqlalchemy.orm import Session
import sys
from pathlib import Path

# Add backend directory to path for imports
backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))

from app.services.ai_service import AIService
from app.db.models import CV, User
from app.utils.logger import logger


class CVGenerator:
    """CV generation service with AI enhancements."""
    
    def __init__(self):
        self.ai_service = AIService()
    
    def generate_cv_with_market_analysis(
        self,
        user_id: int,
        personal_info: Dict[str, Any],
        experience: List[Dict[str, Any]],
        education: List[Dict[str, Any]],
        skills: Dict[str, Any],
        industry: Optional[str] = None,
        awards: List[Dict[str, Any]] = None,
        publications: List[Dict[str, Any]] = None,
        projects: List[Dict[str, Any]] = None,
        memberships: List[Dict[str, Any]] = None,
        job_id: Optional[int] = None,
        photo_url: Optional[str] = None,
        db: Session = None
    ) -> Dict[str, Any]:
        """
        Generate CV with market analysis FIRST, then AI enhancements.
        This ensures suggestions are based on real market trends.
        """
        logger.info(f"Generating CV with market analysis for user {user_id}")
        
        # STEP 1: Analyze job market for the user's industry
        market_analysis = {}
        if db and industry:
            market_analysis = self.ai_service.analyze_job_market(db, sector=industry)
            logger.info(f"Market analysis: {len(market_analysis.get('trending_skills', []))} trending skills identified")
        
        # STEP 2: Generate CV (now informed by market analysis)
        cv_result = self.generate_cv(
            user_id=user_id,
            personal_info=personal_info,
            experience=experience,
            education=education,
            skills=skills,
            awards=awards,
            publications=publications,
            projects=projects,
            memberships=memberships,
            job_id=job_id,
            photo_url=photo_url,
            db=db
        )
        
        # STEP 3: Add market analysis to CV result
        cv_result["market_analysis"] = market_analysis
        
        return cv_result
    
    def generate_cv(
        self,
        user_id: int,
        personal_info: Dict[str, Any],
        experience: List[Dict[str, Any]],
        education: List[Dict[str, Any]],
        skills: Dict[str, Any],
        awards: List[Dict[str, Any]] = None,
        publications: List[Dict[str, Any]] = None,
        projects: List[Dict[str, Any]] = None,
        memberships: List[Dict[str, Any]] = None,
        job_id: Optional[int] = None,
        photo_url: Optional[str] = None,
        db: Session = None
    ) -> Dict[str, Any]:
        """
        Generate a professional CV with AI enhancements.
        
        Args:
            user_id: User ID
            personal_info: Personal information dict
            experience: List of work experience entries
            education: List of education entries
            skills: Skills dictionary (Europass format)
            awards: Optional awards list
            publications: Optional publications list
            projects: Optional projects list
            memberships: Optional memberships list
            job_id: Optional job ID to tailor CV
            photo_url: Optional photo URL
            db: Database session
            
        Returns:
            Generated CV data with AI enhancements
        """
        logger.info(f"Generating CV for user {user_id}")
        
        # Get user
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise ValueError(f"User {user_id} not found")
        
        # Prepare user data
        user_data = {
            "full_name": personal_info.get("first_name", "") + " " + personal_info.get("surname", "") or user.full_name,
            "email": personal_info.get("email") or user.email,
            "wallet_address": user.wallet_address,
            "role": user.role.value,
            "photo_url": photo_url,
            **personal_info
        }
        
        # Generate CV with AI (certificates removed - using education list)
        cv_json = self.ai_service.generate_cv(
            user_data=user_data,
            certificates=[],  # Certificates removed - use education list instead
            experience=experience,
            education=education,
            skills=skills,
            awards=awards or [],
            publications=publications or [],
            projects=projects or [],
            memberships=memberships or [],
            job_id=job_id,
            db=db
        )
        
        # Save CV
        cv = CV(
            user_id=user_id,
            json_content=cv_json,
            ai_score=cv_json.get("ai_score"),
            photo_url=photo_url
        )
        
        db.add(cv)
        db.commit()
        db.refresh(cv)
        
        logger.info(f"CV generated with score: {cv.ai_score}")
        return {
            "id": cv.id,
            "user_id": cv.user_id,
            "json_content": cv.json_content,
            "ai_score": cv.ai_score,
            "photo_url": cv.photo_url,
            "created_at": cv.created_at.isoformat() if cv.created_at else None,
            "updated_at": cv.updated_at.isoformat() if cv.updated_at else None,
        }
    
    def get_cv(self, user_id: int, db: Session) -> Optional[Dict[str, Any]]:
        """Get the latest CV for a user."""
        cv = db.query(CV).filter(CV.user_id == user_id).order_by(CV.created_at.desc()).first()
        if not cv:
            return None
        
        return {
            "id": cv.id,
            "user_id": cv.user_id,
            "json_content": cv.json_content,
            "ai_score": cv.ai_score,
            "photo_url": cv.photo_url,
            "created_at": cv.created_at.isoformat() if cv.created_at else None,
            "updated_at": cv.updated_at.isoformat() if cv.updated_at else None,
        }
    
    def export_to_pdf(self, cv_data: Dict[str, Any]) -> bytes:
        """
        Export CV to PDF format using reportlab.
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib import colors
            from io import BytesIO
            
            logger.info("Exporting CV to PDF")
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=12,
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=6,
                spaceBefore=12,
            )
            
            # Personal Info
            personal_info = cv_data.get("personal_info", {})
            if personal_info:
                name = personal_info.get("full_name", "CV")
                email = personal_info.get("email", "")
                phone = personal_info.get("phone", "")
                address = personal_info.get("address", "")
                
                story.append(Paragraph(name, title_style))
                contact_info = []
                if email:
                    contact_info.append(email)
                if phone:
                    contact_info.append(phone)
                if address:
                    contact_info.append(address)
                if contact_info:
                    story.append(Paragraph(" | ".join(contact_info), styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # Experience
            experience = cv_data.get("experience", [])
            if experience:
                story.append(Paragraph("Experience", heading_style))
                for exp in experience:
                    title = exp.get("job_title", "")
                    company = exp.get("company", "")
                    period = exp.get("period", "")
                    description = exp.get("description", "")
                    
                    exp_text = f"<b>{title}</b>"
                    if company:
                        exp_text += f" - {company}"
                    if period:
                        exp_text += f" ({period})"
                    story.append(Paragraph(exp_text, styles['Normal']))
                    if description:
                        story.append(Paragraph(description, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Education
            education = cv_data.get("education", [])
            if education:
                story.append(Paragraph("Education", heading_style))
                for edu in education:
                    degree = edu.get("degree", "")
                    institution = edu.get("institution", "")
                    period = edu.get("period", "")
                    
                    edu_text = f"<b>{degree}</b>"
                    if institution:
                        edu_text += f" - {institution}"
                    if period:
                        edu_text += f" ({period})"
                    story.append(Paragraph(edu_text, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Skills
            skills = cv_data.get("skills", {})
            if skills:
                story.append(Paragraph("Skills", heading_style))
                skills_text = []
                if isinstance(skills, dict):
                    for category, skill_list in skills.items():
                        if isinstance(skill_list, list):
                            skills_text.append(f"{category}: {', '.join(skill_list)}")
                        else:
                            skills_text.append(f"{category}: {skill_list}")
                else:
                    skills_text.append(str(skills))
                
                if skills_text:
                    story.append(Paragraph(" | ".join(skills_text), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("reportlab not installed. Install with: pip install reportlab")
            raise NotImplementedError("PDF export requires reportlab. Install with: pip install reportlab")
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}", exc_info=True)
            raise
    
    def export_to_word(self, cv_data: Dict[str, Any]) -> bytes:
        """
        Export CV to Word format using python-docx.
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            
            logger.info("Exporting CV to Word")
            
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Personal Info
            personal_info = cv_data.get("personal_info", {})
            if personal_info:
                name = personal_info.get("full_name", "CV")
                email = personal_info.get("email", "")
                phone = personal_info.get("phone", "")
                address = personal_info.get("address", "")
                
                # Name as title
                title = doc.add_heading(name, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Contact info
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_info = []
                if email:
                    contact_info.append(email)
                if phone:
                    contact_info.append(phone)
                if address:
                    contact_info.append(address)
                if contact_info:
                    contact_para.add_run(" | ".join(contact_info))
                
                doc.add_paragraph()  # Spacing
            
            # Experience
            experience = cv_data.get("experience", [])
            if experience:
                doc.add_heading('Experience', level=1)
                for exp in experience:
                    title = exp.get("job_title", "")
                    company = exp.get("company", "")
                    period = exp.get("period", "")
                    description = exp.get("description", "")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{title}").bold = True
                    if company:
                        p.add_run(f" - {company}")
                    if period:
                        p.add_run(f" ({period})")
                    
                    if description:
                        doc.add_paragraph(description, style='List Bullet')
            
            # Education
            education = cv_data.get("education", [])
            if education:
                doc.add_heading('Education', level=1)
                for edu in education:
                    degree = edu.get("degree", "")
                    institution = edu.get("institution", "")
                    period = edu.get("period", "")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{degree}").bold = True
                    if institution:
                        p.add_run(f" - {institution}")
                    if period:
                        p.add_run(f" ({period})")
            
            # Skills
            skills = cv_data.get("skills", {})
            if skills:
                doc.add_heading('Skills', level=1)
                if isinstance(skills, dict):
                    for category, skill_list in skills.items():
                        p = doc.add_paragraph()
                        p.add_run(f"{category}: ").bold = True
                        if isinstance(skill_list, list):
                            p.add_run(", ".join(skill_list))
                        else:
                            p.add_run(str(skill_list))
                else:
                    doc.add_paragraph(str(skills))
            
            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise NotImplementedError("Word export requires python-docx. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}", exc_info=True)
            raise

